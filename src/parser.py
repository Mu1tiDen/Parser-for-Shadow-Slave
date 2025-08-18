from docx import Document
from bs4 import BeautifulSoup
import requests
from sqlalchemy.sql.operators import contains

import database
import os
import re

def put_chapter_in_file_and_jump_to_next(url):
    global path_to_chapters

    headers = {
        "Accept": "*/*",
        "User-Agent": "Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:137.0) Gecko/20100101 Firefox/137.0"
    }

    req = requests.get(url, headers=headers)
    src = req.text
    soup = BeautifulSoup(src, "lxml")

    chapter_title = soup.find(class_="tl_article_content").find("h1")
    chapter_text = soup.find(class_="tl_article_content").find_all("p")

    head_number =  re.sub(':.*', '', chapter_title.text)
    path = path_to_chapters + "/" + head_number

    if not (os.path.exists(path) and os.path.isdir(path)):
        os.mkdir(path)

    txt_path = path + "/" + head_number + ".txt"
    docx_path = path + "/" + head_number + ".docx"

    if os.path.exists(docx_path): os.remove(docx_path)
    if os.path.exists(txt_path): os.remove(txt_path)


    with open(txt_path, "a") as file:
        file.write(f'{chapter_title.text.upper()}\n')
        for text_line in chapter_text:
            filtered_line = text_line.text
            if filtered_line != '\n' and filtered_line not in "Предыдущая глава Следующая глава":
                file.write(f'{filtered_line}\n')
        file.write('\n')

    document = Document()
    title = document.add_heading(chapter_title.text, 1)
    title.bold = True
    for text_line in chapter_text:
        filtered_line = text_line.text
        if filtered_line != '\n' and filtered_line not in "Предыдущая глава Следующая глава":
            document.add_paragraph(filtered_line, style="No Spacing")
    document.save(docx_path)

    database.add_chapter_in_db(head_number[5:], re.sub(".*: ", '', chapter_title.text), url, path)
    print(f'"{chapter_title.text}" recorded successfully!\n')

    links = soup.find_all("a", href=True)

    for link in links:
        if link.text in "Следующая глава":
            #print("New link successfully write!")
            next_link = f'https://telegra.ph{link.get('href')}'
            return next_link.rstrip("#,;/")
    return "NULL"


path_to_chapters = 'data/chapters'
if not (os.path.exists(path_to_chapters) and os.path.isdir(path_to_chapters)):
    os.mkdir(path_to_chapters)

url = 'https://telegra.ph/Glava-2261-Vypolnenie-obeshchaniya-04-12'
while True:
    url = put_chapter_in_file_and_jump_to_next(url)

    if url == 'NULL':
        break

